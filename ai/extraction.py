# user_masters/ai/extraction.py
import io
import time
import json
import logging
import pandas as pd
import streamlit as st
import fitz
import docx
from openai import OpenAI
from concurrent.futures import ThreadPoolExecutor, as_completed
import sys
import os
import re

# Prevent Name Shadowing: If run directly, the current file 'extraction.py' conflicts 
# with the 'extraction' package. We clean the path and modules dynamically.
if __name__ == "__main__" or sys.path[0].endswith("ai"):
    parent_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    current_dir = os.path.dirname(os.path.abspath(__file__))
    if parent_dir not in sys.path:
        sys.path.insert(0, parent_dir)
    if current_dir in sys.path:
        sys.path.remove(current_dir)
    if 'extraction' in sys.modules and sys.modules['extraction'].__file__ == __file__:
        del sys.modules['extraction']

from config.constants import (
    USER_MASTER_COLS, TICK_VALUES, ROLE_NEGATIVE_VALUES,
    MAX_FILE_SIZE_MB, MAX_PDF_PAGES, MAX_AI_CONTEXT_KB,
    AI_RETRY_ATTEMPTS, AI_RETRY_BASE_WAIT, AI_ALLOWED_EDIT_COLS
)
from models.schemas import UserMasterResult, AISmartResponse, RowUpdate, VerificationResult
from models.dataframe_contract import enforce_contract
from extraction.merge import _merge_duplicate_users
from extraction.utils import (
    get_all_api_keys, find_matching_excel_roles, filter_sheets_by_intent,
    detect_header_row, check_is_sub_header, build_unique_headers, detect_tick_role_columns,
    build_temp_col_mapping
)
from utils.common import is_empty_value, has_value

log = logging.getLogger(__name__)

USER_EXTRACTION_PROMPT = """
You are a High-Precision, Format-Agnostic User Data Extraction Engine.
Your task is to extract a structured user list from ANY document format — staff rosters, HR matrices, Word tables, PDFs, CSV exports, or any custom layout.

=== STEP 1: UNDERSTAND THE STRUCTURE ===
Before extracting, ANALYZE the document layout:
- Is this a STANDARD ROSTER? (one row per person with named columns like Name, Email, Department...)
- Is this a ROLE MATRIX? (columns represent roles, cells are Yes/No, and a separate column has the person name)
- Is this a FREE-FORMAT LIST? (names and details listed without a strict table structure)
- Could there be MERGED HEADERS or multi-row headers above the data?
DO NOT assume any fixed column order or name. Identify each column by its CONTENT and SEMANTIC MEANING.

=== STEP 2: FIELD IDENTIFICATION ===
Identify fields by MEANING, not by position or exact column header name:
- NAME field: Contains full person names including titles (e.g. "Dr. NG SWEET MAY", "Mr. RAMESH JOSHI"). Capture prefixes as part of the first name.
- EMAIL field: Contains @ symbols
- PHONE/MOBILE: Contains 010/011/012 patterns or dashes like 012-345-6789
- EMPLOYEE ID: Alphanumeric codes like CA00001, GB10007, EMP001
- DEPARTMENT: Hospital departments (Emergency Room, ICU, OPD, Surgery...)
- UNIT/FACILITY: Hospital branches or facility names (CAH Bukit Jalil, Ankura...)
- DESIGNATION/POSITION: Job titles (Doctor, Nurse, Consultant, Locum MO...)
- ROLES: Either a column with slash-separated text OR multiple Yes/No columns whose HEADERS are role names

=== STEP 3: EXTRACTION RULES ===
1. ROW INTEGRITY (MANDATORY): You MUST extract EVERY person in the data. Do not skip anyone. Do not summarize.
2. MULTI-USER SPLIT RULE (MANDATORY): A single row may contain data for MULTIPLE people, separated by a CONSISTENT DELIMITER. The delimiter can be ANY of: "|", ",", ";", "/", " & ", " and " — whichever is used consistently across that row.
   DETECTION: If the SAME delimiter appears in MULTIPLE columns of the same row (e.g. firstName has "Bijay | Maumita" AND employeeId has "1035708 | 192004"), that row contains multiple users. You MUST create SEPARATE JSON objects for each person.
   SPLITTING RULES:
   - Identify the delimiter being used (e.g. "|" or "," or "&").
   - User 1 takes the 1st part of EVERY delimited column.
   - User 2 takes the 2nd part of EVERY delimited column.
   - For columns with NO delimiter (e.g. middleName "Krishna" when others use "|"), apply the value ONLY to User 1. Leave it BLANK for User 2.
   - NEVER include the delimiter symbol in your final JSON fields — strip it out completely.
   - WORKED EXAMPLE with "|" delimiter (follow exactly):
       Source row: firstName="Bijay | Maumita", middleName="Krishna", lastName="Maity | Saha", employeeId="1035708 | 192004"
       firstName has "|" → 2 users. lastName has "|" → 2 users. employeeId has "|" → 2 users. middleName has NO "|" → belongs to User 1 only.
       Correct output:
         User 1: firstName="Bijay",   middleName="Krishna", lastName="Maity", employeeId="1035708"
         User 2: firstName="Maumita", middleName="",        lastName="Saha",  employeeId="192004"
   - WORKED EXAMPLE with "|" delimiter and MISSING lastName/employeeId for User 2 (follow exactly):
        Source row: firstName="Arindam | Riya", middleName="", lastName="Chakraborty", employeeId="10092"
        firstName has "|" → 2 users. middleName has NO "|" → blank for both. lastName has NO "|" → belongs ONLY to User 1. employeeId has NO "|" → belongs ONLY to User 1.
        Correct output:
          User 1: firstName="Arindam", middleName="", lastName="Chakraborty", employeeId="10092"
          User 2: firstName="Riya",    middleName="", lastName="",            employeeId=""
        WRONG output (never do this):
          User 2: firstName="Riya", lastName="Chakraborty"  ← copying no-pipe lastName to User 2 is FORBIDDEN.
          User 2: employeeId="10092" ← copying no-pipe employeeId to User 2 is FORBIDDEN.
   - WORKED EXAMPLE with 3-user split and NO middle names (follow exactly):
       Source row: firstName="Priyadarshini | Manisha | Anasuya", middleName="", lastName="Roy | Santra | Bajpaye"
       Correct output:
         User 1: firstName="Priyadarshini", lastName="Roy"
         User 2: firstName="Manisha",      lastName="Santra"
         User 3: firstName="Anasuya",     lastName="Bajpaye"
       WRONG output (never do this):
         User 1: firstName="Priyadarshini Manisha", lastName="Roy"  ← Merging names is FORBIDDEN.
   - WORKED EXAMPLE with "," delimiter:
       Source row: firstName="Anjali, Priya", lastName="Sen, Sharma", employeeId="101, 102"
       Correct output:
         User 1: firstName="Anjali", lastName="Sen",    employeeId="101"
         User 2: firstName="Priya",  lastName="Sharma", employeeId="102"
 3. FIELD MAPPING:
    - firstName, middleName, lastName: Extract logically from the source data.
    - STRICT COLUMN ALIGNMENT (CRITICAL): Do NOT shift data across columns. If the "Middle Name" column in Excel is blank for a specific user, the "middleName" field in JSON MUST be blank. NEVER move the second person's first name (from a pipe-split) into the first person's middleName field.
    - If the document has a "Full Name" column, split it logically into firstName and lastName.
    - employeeId: Extract the unique ID.
    - roles: This is CRITICAL. In this Excel format, the roles are often found in Column A as "floating headers" (e.g. 'Audit User - CESC - CAUTI') or are provided in the [ROLE SECTION:] tag at the start of the row. Combine all applicable roles found in Column A and [ROLE SECTION:] into this field, separated by "|".
    - departments / units / designation / email / phone: Map correctly based on column headers.
4. UNIT NAME (STRICT RULE): Do NOT guess, assume, or infer the unit name from email addresses, department names, or any other fields. The `units` field must remain completely empty unless the Excel sheet has an explicit column for the unit name and the cells under it explicitly contain the unit name value.
5. NO LOGIC: Do not try to generate usernames or complex passwords. Just extract the raw name parts. Python will handle the rest.
6. NA LOGIC: If a value is "-", leave it blank.
7. ROLE SECTION TAGS (CRITICAL): Lines tagged with [ROLE SECTION: ...] are SECTION HEADERS, NOT people. NEVER create a user record from a [ROLE SECTION:] tag. Only use them to determine the `roles` value for the real data rows that follow. If a row has no name, no email, and no employee ID, SKIP it entirely. Always extract rows that contain an email address or active cells, even if they have a non-standard name (like "MMC Rehab") or lack an employee ID.
8. MIDDLE NAME RULE (CRITICAL): The `middleName` field is ONLY for a person's own middle name or initial (e.g. "K", "Kumar", "Rani"). It must NEVER contain another person's first name from a pipe-split. If firstName is "Arindam | Riya", you must create TWO separate users: User 1 has firstName="Arindam" and middleName="" (blank). User 2 has firstName="Riya" and middleName="" (blank). The text after the "|" is a second person, not a middle name.

OUTPUT FORMAT (JSON):
{{
  "document_name": "",
  "users": [
    {{
      "userName": "",
      "firstName": "",
      "middleName": "",
      "lastName": "",
      "employeeId": "",
      "departments": "",
      "roles": "",
      "units": "",
      "designation": "",
      "email": "",
      "phone": "",
      "isEnabled": "Yes"
    }}
  ]
}}
"""

def get_openai_client(api_key):
    if not api_key:
        return None
    api_key_str = str(api_key).strip()
    if api_key_str.startswith("AIzaSy"):
        return OpenAI(
            api_key=api_key_str,
            base_url="https://generativelanguage.googleapis.com/v1beta/openai/"
        )
    return OpenAI(api_key=api_key_str)

def probe_api_key(key):
    """Validates an API key using the free /models list endpoint (no quota consumed)."""
    if not key:
        return False
    key_str = str(key).strip()
    client = get_openai_client(key_str)
    if not client:
        return False
    try:
        # models.list() is a free REST call — validates key authorization without spending tokens
        client.models.list()
        return True
    except Exception as e:
        log.warning("Pre-flight check failed for key prefix %s...: %s: %s", key_str[:12], type(e).__name__, e)
        return False

def get_healthy_api_keys(api_key):
    """
    Returns only verified, active API keys.
    Results are cached in Streamlit session state with a 30-minute TTL to avoid
    running probes on every Streamlit rerun within the same session.
    """
    all_keys = get_all_api_keys(api_key)
    if not all_keys:
        return []

    _CACHE_TTL_SECONDS = 1800  # 30 minutes
    cached = st.session_state.get("healthy_api_keys")
    cached_ts = st.session_state.get("healthy_api_keys_ts", 0)
    if cached is not None and (time.time() - cached_ts) < _CACHE_TTL_SECONDS:
        still_valid = [k for k in cached if k in all_keys]
        if still_valid:
            return still_valid

    healthy_keys = []
    with ThreadPoolExecutor(max_workers=len(all_keys)) as executor:
        future_to_key = {executor.submit(probe_api_key, key): key for key in all_keys}
        for future in as_completed(future_to_key):
            key = future_to_key[future]
            try:
                if future.result():
                    healthy_keys.append(key)
            except Exception:
                pass

    ordered_healthy = [k for k in all_keys if k in healthy_keys]
    st.session_state.healthy_api_keys = ordered_healthy
    st.session_state.healthy_api_keys_ts = time.time()
    return ordered_healthy

def verify_extracted_user_source(user_dict, raw_chunk_text):
    """
    Deterministically verifies that all key fields (firstName, lastName, email, employeeId, userName)
    extracted by the LLM are valid substrings in the raw source chunk text.
    Also enforces that the record contains at least one identifier (Person Validation Rule).
    """
    def _clean_val(val):
        if val is None:
            return ""
        s = str(val).strip()
        if s.lower() in ('nan', 'none', '-', ''):
            return ""
        return s

    first_name = _clean_val(user_dict.get('firstName'))
    email = _clean_val(user_dict.get('email'))
    emp_id = _clean_val(user_dict.get('employeeId'))
    username = _clean_val(user_dict.get('userName'))
    last_name = _clean_val(user_dict.get('lastName'))
    
    # 1. Person Validation Rule: must contain at least one valid identifier
    if not (first_name or email or emp_id or username):
        log.warning("Discarding record: fails Person Validation Rule (missing firstName, email, employeeId, and userName)")
        return False

    # Normalize source text for substring searches
    source_lower = str(raw_chunk_text).lower()

    # Helper function to check if string exists in source
    def _is_in_source(val_clean):
        if not val_clean:
            return True
        # Strip potential trailing/leading whitespace or quote characters
        val_clean = val_clean.strip('"').strip("'").lower()
        if not val_clean:
            return True
        return val_clean in source_lower

    # Check firstName
    if not _is_in_source(first_name):
        log.warning("Hallucination detected: firstName '%s' not in source text.", first_name)
        return False

    # Check lastName
    if not _is_in_source(last_name):
        log.warning("Hallucination detected: lastName '%s' not in source text.", last_name)
        return False

    # Check employeeId
    if not _is_in_source(emp_id):
        log.warning("Hallucination detected: employeeId '%s' not in source text.", emp_id)
        return False

    # Check email
    if not _is_in_source(email):
        log.warning("Hallucination detected: email '%s' not in source text.", email)
        return False

    # Check userName
    if not _is_in_source(username):
        log.warning("Hallucination detected: userName '%s' not in source text.", username)
        return False

    return True


def cross_examine_extracted_users(client, model, raw_chunk_text, extracted_users):
    """
    Performs dual-model cross-examination to verify the extracted users
    list against the raw text chunk using a Structured Output verification prompt.
    """
    if not extracted_users:
        return True

    # Limit verification checks to prevent infinite loops
    try:
        users_json = json.dumps([u for u in extracted_users], default=str)
        prompt = f"""
        You are an Independent Data Auditor.
        Below is a raw source data chunk followed by a list of users extracted from it.

        RAW SOURCE DATA:
        \"\"\"{raw_chunk_text}\"\"\"

        EXTRACTED USERS (JSON):
        {users_json}

        TASK:
        1. Compare every single field of the extracted users list with the raw source data.
        2. Set `is_hallucinated` to True ONLY if there are fabricated records, names, emails, or employee IDs that do NOT exist anywhere in the RAW SOURCE DATA.
        3. Do NOT flag minor casing differences or spacing differences as hallucinations. Only flag completely fabricated or made-up details.
        """

        completion = client.chat.completions.parse(
            model=model,
            messages=[
                {"role": "system", "content": "You are a precise data auditor."},
                {"role": "user", "content": prompt}
            ],
            response_format=VerificationResult,
            timeout=30,
            temperature=0.0
        )
        res = completion.choices[0].message.parsed
        if res and res.is_hallucinated:
            log.warning("Dual-Model cross-examination flagged batch as hallucinated. Reason: %s", res.reason)
            return False
    except Exception as e:
        log.warning("An error occurred during dual-model cross-examination: %s", e)
        # Fallback to True to avoid dropping data on network/timeout issues
        return True
    return True


def openai_extract_users(file_bytes, filename, api_key, intent="", pass_prefix="Med"):
    """Universal AI User Extraction Engine with Chunking & Failover."""
    # ── File-size guard ───────────────────────────────────────────────────────
    size_mb = len(file_bytes) / 1_048_576
    if size_mb > MAX_FILE_SIZE_MB:
        st.error(f"❌ **{filename}** is {size_mb:.1f} MB — exceeds the {MAX_FILE_SIZE_MB} MB limit. "
                 f"Please split the file and re-upload.")
        return pd.DataFrame()
    # ─────────────────────────────────────────────────────────────────────────

    try:
        healthy_keys = get_healthy_api_keys(api_key)
        if not healthy_keys:
            log.warning("No healthy API keys available. Skipping AI extraction.")
            return None
            
        # Scaled OpenAI workers limit
        max_workers = min(len(healthy_keys) * 5, 5)
        log.info("Starting AI extraction. Healthy keys: %d, Workers: %d", len(healthy_keys), max_workers)

        dynamic_prompt = USER_EXTRACTION_PROMPT
        ext = filename.lower()
        
        if ext.endswith(('.xlsx', '.xls', '.csv')):
            if ext.endswith('.csv'):
                raw_df = pd.read_csv(io.BytesIO(file_bytes), header=None)
                sheet_dfs = {'Sheet1': raw_df}
            else:
                # Read ALL sheets — sheet_name=None returns {sheet_name: DataFrame}
                all_sheets = pd.read_excel(io.BytesIO(file_bytes), header=None, sheet_name=None)
                sheet_dfs = filter_sheets_by_intent(all_sheets, intent)
                st.toast(f"📋 Found {len(sheet_dfs)} sheet(s): {', '.join(sheet_dfs.keys())}")

            # We collect all chunks across all sheets, each with its own correct header context
            all_chunks_to_process = []
            has_tick_role_columns = False
            global_excel_rows_data = []

            for sheet_name, raw_df in sheet_dfs.items():
                raw_df = raw_df.dropna(how='all').reset_index(drop=True)
                mask = raw_df.astype(str).apply(lambda x: x.str.contains(r'[a-zA-Z0-9]', na=False)).any(axis=1)
                raw_df = raw_df[mask].reset_index(drop=True)
                if raw_df.empty:
                    continue

                # --- Auto-detect header and sub-header row (aligned with local_extract_users) ---
                header_row_idx = detect_header_row(raw_df)

                headers_raw = [str(h).strip() for h in raw_df.iloc[header_row_idx].values]
                col_mapping_temp = build_temp_col_mapping(headers_raw)
                is_sub_header = check_is_sub_header(raw_df, header_row_idx, col_mapping_temp)
                headers = build_unique_headers(raw_df, header_row_idx, is_sub_header)
                first_data_row = header_row_idx + 2 if is_sub_header else header_row_idx + 1
                header_rows_df = raw_df.iloc[:first_data_row]
                data_rows_df = raw_df.iloc[first_data_row:]

                def row_to_str(row):
                    return " ; ".join([str(v).strip().replace(';', ',') if str(v).strip().lower() not in ('nan', 'none', '') else '-' for v in row.values])

                sheet_header_context = f"COLUMN HEADERS (from '{sheet_name}'):\n" + "\n".join(row_to_str(r) for _, r in header_rows_df.iterrows())

                role_cols = detect_tick_role_columns(headers, data_rows_df)
                if role_cols:
                    has_tick_role_columns = True

                # --- EXTRACT ROLE TICKS PER ROW ---
                for idx, row in data_rows_df.iterrows():
                    row_roles = []
                    for col_idx, role_name in role_cols.items():
                        val = str(row.iloc[col_idx]).strip().lower() if col_idx < len(row) else ""

                        is_ticked = False
                        if 'module|' in role_name.lower():
                            is_ticked = val not in ROLE_NEGATIVE_VALUES
                        else:
                            is_ticked = val in TICK_VALUES

                        if is_ticked:
                            clean_role_name = role_name
                            if '|' in clean_role_name:
                                clean_role_name = clean_role_name.split('|')[-1]
                            clean_role_name = re.sub(r'_\d+$', '', clean_role_name)
                            row_roles.append(clean_role_name)

                    raw_vals = []
                    for v in row.values:
                        if pd.notna(v):
                            v_str = str(v).strip().lower()
                            if has_value(v_str):
                                raw_vals.append(v_str)

                    global_excel_rows_data.append({
                        'roles': '|'.join(row_roles),
                        'raw_values': raw_vals
                    })

                # --- DYNAMIC FORMAT CLASSIFICATION FOR COLUMN A ---
                col_a_values = data_rows_df.iloc[:, 0].dropna().astype(str).str.strip() if len(data_rows_df) > 0 else pd.Series()
                col_a_non_empty = col_a_values[~col_a_values.str.lower().isin(['nan', 'none', '', '-'])]
                non_empty_ratio = len(col_a_non_empty) / len(data_rows_df) if len(data_rows_df) > 0 else 0

                col_a_header = str(headers[0]).lower().strip() if len(headers) > 0 else ""
                is_col_a_user_attr = any(kw in col_a_header for kw in ['name', 'username', 'employee', 'emp', 'id', 'email', 'mail', 'phone', 'mobile'])

                is_floating_role_section = (non_empty_ratio < 0.45) and (not is_col_a_user_attr)
                log.info("Sheet '%s' classification: non_empty_ratio=%.2f, is_col_a_user_attr=%s -> is_floating_role_section=%s", sheet_name, non_empty_ratio, is_col_a_user_attr, is_floating_role_section)

                sheet_lines = []
                last_col_a_value = ""
                for _, row in data_rows_df.iterrows():
                    line = row_to_str(row)
                    if not line.strip():
                        continue
                    col_a_val = str(row.iloc[0]).strip() if len(row) > 0 else ""
                    if has_value(col_a_val):
                        last_col_a_value = col_a_val
                    tag = f"[SHEET: {sheet_name}]"
                    if is_floating_role_section and last_col_a_value:
                        tag += f" [ROLE SECTION: {last_col_a_value}]"
                    sheet_lines.append(f"{tag} {line}")

                # Small chunks (20 rows) = higher precision, fewer skipped users, no API timeouts
                chunk_size = 20
                for i in range(0, len(sheet_lines), chunk_size):
                    chunk = sheet_lines[i:i + chunk_size]
                    all_chunks_to_process.append({
                        'text': sheet_header_context + "\n\nDATA:\n" + "\n".join(chunk),
                        'sheet': sheet_name,
                        'line_start': i
                    })

            if not all_chunks_to_process:
                st.error("⚠️ No data found in any of the Excel sheets.")
                return pd.DataFrame()
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            status_text.info(f"🚀 Started extraction of {len(all_chunks_to_process)} batches...")
            
            all_users = []
            
            def process_chunk(chunk_data, batch_idx):
                chunk_text = chunk_data['text']
                # Round-robin: assign each batch a primary key to spread load
                key_order = [healthy_keys[(batch_idx + i) % len(healthy_keys)] for i in range(len(healthy_keys))]
                
                for k_idx, current_key in enumerate(key_order):
                    current_client = get_openai_client(current_key)
                    if not current_client:
                        continue
                    
                    if current_key.startswith("AIzaSy"):
                        models_to_try = ["gemini-1.5-flash", "gemini-1.5-pro"]
                    else:
                        models_to_try = ["gpt-4o-mini", "gpt-4o"]
                        
                    for model in models_to_try:
                        try:
                            completion = current_client.chat.completions.parse(
                                model=model,
                                messages=[
                                    {"role": "system", "content": dynamic_prompt},
                                    {"role": "user", "content": f"USER INTENT: {intent}\n\nINPUT DATA TABLE:\n{chunk_text}"}
                                ],
                                response_format=UserMasterResult,
                                timeout=60,
                                temperature=0.0
                            )
                            result = [u.__dict__ for u in completion.choices[0].message.parsed.users]
                            if not result and model in ("gpt-4o-mini", "gemini-1.5-flash"):
                                time.sleep(1)
                                completion2 = current_client.chat.completions.parse(
                                    model=model,
                                    messages=[
                                        {"role": "system", "content": dynamic_prompt},
                                        {"role": "user", "content": f"USER INTENT: {intent}\n\nINPUT DATA TABLE:\n{chunk_text}"}
                                    ],
                                    response_format=UserMasterResult,
                                    timeout=60,
                                    temperature=0.0
                                )
                                result = [u.__dict__ for u in completion2.choices[0].message.parsed.users]
                            
                            # Verify and cross-examine
                            verified = [u for u in result if verify_extracted_user_source(u, chunk_text)]
                            if verified:
                                if not cross_examine_extracted_users(current_client, model, chunk_text, verified):
                                    log.warning("Batch %d failed cross-examination. Keeping users to prevent dropping data.", batch_idx)
                            return verified
                        except Exception as e:
                            log.warning("Batch %d AI error model=%s: %s", batch_idx, model, e, exc_info=True)
                            err_str = str(e).lower()
                            if "quota" in err_str or "rate limit" in err_str or "429" in err_str:
                                log.info("Batch %d key #%d rate limited. Trying next key...", batch_idx, k_idx + 1)
                                break
                            else:
                                if model == models_to_try[0] and len(models_to_try) > 1:
                                    time.sleep(1)
                                    continue
                return []

            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                future_to_idx = {executor.submit(process_chunk, chunk, i): i for i, chunk in enumerate(all_chunks_to_process)}
                completed = 0
                # Store results by index to preserve source order
                ordered_results = {}
                for future in as_completed(future_to_idx):
                    idx = future_to_idx[future]
                    ordered_results[idx] = future.result()
                    completed += 1
                    progress_bar.progress(completed / len(all_chunks_to_process))
                    captured = sum(len(v) for v in ordered_results.values())
                    status_text.info(f"⚡ Extraction: {completed}/{len(all_chunks_to_process)} batches done ({captured} users captured)")
            
            # Flatten in chunk order to preserve original row sequence
            all_users = []
            for i in sorted(ordered_results.keys()):
                all_users.extend(ordered_results[i])
            
            status_text.empty()
            progress_bar.empty()
            
            # Override roles with exact Excel ticks if available
            if has_tick_role_columns and global_excel_rows_data:
                for user in all_users:
                    matched_roles = find_matching_excel_roles(user, global_excel_rows_data)
                    if matched_roles:
                        user['roles'] = matched_roles
            
            raw_df = pd.DataFrame(all_users)
            result_df = _merge_duplicate_users(raw_df, pass_prefix=pass_prefix)
            result_df = enforce_contract(result_df)
            st.toast(f"✅ {len(raw_df)} raw rows → {len(result_df)} unique users after merge.")
            return result_df

        elif ext.endswith('.pdf'):
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            total_pages = len(doc)
            if total_pages > MAX_PDF_PAGES:
                st.warning(f"⚠️ **{filename}** has {total_pages} pages — processing first {MAX_PDF_PAGES} only.")
            all_lines = []
            for page in doc.pages(0, min(total_pages, MAX_PDF_PAGES)):
                text = page.get_text("text")
                lines = [line.strip() for line in text.split("\n") if line.strip()]
                all_lines.extend(lines)
            st.toast(f"📄 Read {min(total_pages, MAX_PDF_PAGES)}/{total_pages} PDF pages.")
            header_context = "SOURCE: PDF Document\n"
            
        elif ext.endswith(('.docx', '.doc')):
            doc = docx.Document(io.BytesIO(file_bytes))
            all_lines = []
            for para in doc.paragraphs:
                if para.text.strip():
                    all_lines.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    if any(row_data):
                        all_lines.append(" | ".join(row_data))
            st.toast(f"📄 Read Word document.")
            header_context = "SOURCE: Word Document\n"
            
        else:
            st.error(f"Unsupported file format: {ext}")
            return None

        # Fallback for PDF/Word
        chunk_size = 30
        chunks = []
        for i in range(0, len(all_lines), chunk_size):
            chunk = all_lines[i:i + chunk_size]
            chunks.append(header_context + "\n\nDATA:\n" + "\n".join(chunk))
        
        if not chunks: return pd.DataFrame()
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        def process_simple_chunk(chunk_text, batch_idx):
            key_order = [healthy_keys[(batch_idx + i) % len(healthy_keys)] for i in range(len(healthy_keys))]
            for k_idx, current_key in enumerate(key_order):
                current_client = get_openai_client(current_key)
                if not current_client:
                    continue
                
                if current_key.startswith("AIzaSy"):
                    models_to_try = ["gemini-1.5-flash", "gemini-1.5-pro"]
                else:
                    models_to_try = ["gpt-4o-mini", "gpt-4o"]
                    
                for model in models_to_try:
                    try:
                        completion = current_client.chat.completions.parse(
                            model=model,
                            messages=[
                                {"role": "system", "content": dynamic_prompt},
                                {"role": "user", "content": f"USER INTENT: {intent}\n\nINPUT DATA TABLE:\n{chunk_text}"}
                            ],
                            response_format=UserMasterResult,
                            timeout=60,
                            temperature=0.0
                        )
                        result = [u.__dict__ for u in completion.choices[0].message.parsed.users]
                        # Verify and cross-examine
                        verified = [u for u in result if verify_extracted_user_source(u, chunk_text)]
                        if verified:
                            if not cross_examine_extracted_users(current_client, model, chunk_text, verified):
                                log.warning("Simple Batch %d failed cross-examination. Keeping users to prevent dropping data.", batch_idx)
                        return verified
                    except Exception as e:
                        log.warning("PDF/Word Batch %d AI error model=%s: %s", batch_idx, model, e, exc_info=True)
                        err_str = str(e).lower()
                        if "quota" in err_str or "rate limit" in err_str or "429" in err_str:
                            break
                        else:
                            if model == models_to_try[0] and len(models_to_try) > 1:
                                time.sleep(1)
                                continue
            return []

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_idx = {executor.submit(process_simple_chunk, text, i): i for i, text in enumerate(chunks)}
            completed = 0
            ordered_results = {}
            for future in as_completed(future_to_idx):
                idx = future_to_idx[future]
                ordered_results[idx] = future.result()
                completed += 1
                progress_bar.progress(completed / len(chunks))
        
        status_text.empty()
        progress_bar.empty()
        
        all_users = []
        for i in sorted(ordered_results.keys()):
            all_users.extend(ordered_results[i])
        
        raw_df = pd.DataFrame(all_users)
        result_df = _merge_duplicate_users(raw_df, pass_prefix=pass_prefix)
        return enforce_contract(result_df)

    except Exception as e:
        log.error("Exception in openai_extract_users", exc_info=True)
        st.error(f"AI Extraction Error: {str(e)[:200]}")
        return None

def apply_ai_smart_context(df, command, api_key, context_df=None):
    """
    Applies natural language commands to the user dataframe using AI.
    Includes: 60s timeout, exponential-backoff retry, column allowlist guard.
    Optionally accepts a context_df for external lookups/mappings.
    """
    healthy_keys = get_healthy_api_keys(api_key)
    if not healthy_keys:
        return None, "No healthy API keys available. Please configure API keys."
    model = "gpt-4o-mini"

    # Context size guard: truncate rows if JSON would be too large
    context_data_df = df
    context_json = context_data_df.to_json(orient='records')
    _is_truncated = False
    if len(context_json) > MAX_AI_CONTEXT_KB * 1024:
        # Send only first 200 rows for initial context evaluation
        context_data_df = df.head(200)
        context_json = context_data_df.to_json(orient='records')
        _is_truncated = True
        
    mapping_prompt = ""
    if context_df is not None and not context_df.empty:
        mapping_prompt = f"\n\nEXTERNAL MAPPING DATA (UPLOADED FILE) COLUMNS: {list(context_df.columns)}\n"

    def _build_prompt(data_json):
        """Build a complete AI prompt for a given chunk of data JSON."""
        return f"""
    You are a User Master Data Expert. The user wants to modify the following staff list.

    USER COMMAND: {command}
    {mapping_prompt}

    CURRENT DATA (JSON):
    {data_json}

    INSTRUCTIONS & STRUCTURAL MAPPING:
    1. Determine the best way to execute the USER COMMAND.
    2. If the user wants to MAP a column using the EXTERNAL MAPPING DATA, populate ONLY `mapping_intent` in the schema.
    3. If the user wants to perform a SIMPLE FIND AND REPLACE of a text/substring in a column across the rows, populate ONLY `replace_intent` in the schema.
       - IMPORTANT: Use this intent for all search and replace requests. Do NOT perform the replacements manually in the `updates` list if they can be expressed as a find and replace.
    4. If the user wants to SET a column to a fixed value for ALL rows (or rows where another column equals a specific value), populate ONLY `set_value_intent` in the schema.
    5. For all other custom or row-specific edits, populate the `updates` list in the schema:
       - Each update MUST contain the row's '#' (serial number) and ONLY the fields that actually changed.
       - CRITICAL FOR DELIMITED STRINGS (e.g. roles / departments): If you are modifying a pipe-separated (|) list of roles or departments, you MUST preserve all other values in the list. DO NOT truncate or drop other values in the list. For example, if the original value is "Admin|Doctor|User" and you want to replace "User" with "Guest", the updated value MUST be "Admin|Doctor|Guest". Dropping "Admin|Doctor" is a severe failure.
       - DO NOT take shortcuts. You MUST exhaustively process every single matching row.
    """

    prompt = _build_prompt(context_json)

    last_error = "Unknown error"
    for attempt in range(AI_RETRY_ATTEMPTS):
        current_key = healthy_keys[attempt % len(healthy_keys)]
        client = get_openai_client(current_key)
        model = "gemini-1.5-flash" if current_key.startswith("AIzaSy") else "gpt-4o-mini"
        try:
            completion = client.chat.completions.parse(
                model=model,
                messages=[
                    {"role": "system", "content": "You are a data patching engine."},
                    {"role": "user", "content": prompt}
                ],
                response_format=AISmartResponse,
                timeout=120,
                temperature=0.0
            )

            res_data = completion.choices[0].message.parsed
            if not res_data:
                last_error = "AI failed to parse response into Structured Output schema."
                time.sleep(AI_RETRY_BASE_WAIT * (2 ** attempt))
                continue

            # Programmatic Mapping Mode
            if res_data.mapping_intent is not None and context_df is not None:
                intent = res_data.mapping_intent
                target_col = intent.target_col
                lookup_col = intent.lookup_col
                value_col = intent.value_col
                
                if target_col in df.columns and lookup_col in context_df.columns and value_col in context_df.columns:
                    # Create case-insensitive lookup dict
                    lookup_dict = {str(k).strip().lower(): v for k, v in zip(context_df[lookup_col], context_df[value_col]) if pd.notna(k)}
                    
                    new_df = df.copy()
                    affected = 0
                    for i, row in new_df.iterrows():
                        old_val = str(row.get(target_col, '')).strip()
                        if old_val.lower() in lookup_dict:
                            new_val = lookup_dict[old_val.lower()]
                            # If mapped value is empty/NaN, keep the original value untouched
                            if pd.notna(new_val) and str(new_val).strip() != '' and str(new_val).strip().lower() not in ('nan', 'none', '-', 'na', 'n/a'):
                                if new_val != row.get(target_col):
                                    new_df.at[i, target_col] = new_val
                                    affected += 1
                    
                    summary = f"AI programmatically mapped {affected} row(s) in column '{target_col}' using '{lookup_col}' -> '{value_col}'."
                    return new_df, summary
                else:
                    last_error = f"AI generated invalid mapping columns: {target_col}, {lookup_col}, {value_col}"
                    break

            # Programmatic Replace Mode
            if res_data.replace_intent is not None:
                intent = res_data.replace_intent
                target_col = intent.target_col
                search_text = intent.search_text
                replace_text = intent.replace_text
                
                if target_col in df.columns and search_text:
                    new_df = df.copy()
                    escaped_search = re.escape(search_text)
                    
                    # Store original values to count affected rows
                    orig_vals = new_df[target_col].copy()
                    
                    # Convert to string and do regex replacement
                    new_df[target_col] = new_df[target_col].astype(str).str.replace(escaped_search, replace_text, regex=True)
                    new_df[target_col] = new_df[target_col].replace('nan', '')
                    
                    # Calculate affected
                    affected = (orig_vals.astype(str).replace('nan', '') != new_df[target_col]).sum()
                    
                    summary = f"AI programmatically replaced '{search_text}' with '{replace_text}' in {affected} row(s) of '{target_col}'."
                    return new_df, summary

            # Programmatic Set-Value Mode
            if res_data.set_value_intent is not None:
                intent = res_data.set_value_intent
                target_col = intent.target_col
                new_value = intent.value
                filter_col = intent.filter_col
                filter_value = intent.filter_value

                if target_col and target_col in AI_ALLOWED_EDIT_COLS and target_col in df.columns:
                    new_df = df.copy()
                    if filter_col and filter_value is not None and filter_col in df.columns:
                        # Filtered set: only rows where filter_col matches filter_value (case-insensitive)
                        mask = new_df[filter_col].astype(str).str.strip().str.lower() == str(filter_value).strip().lower()
                        affected = int(mask.sum())
                        new_df.loc[mask, target_col] = new_value
                        summary = f"AI locally set '{target_col}' = '{new_value}' for {affected} row(s) where {filter_col} = '{filter_value}'."
                    else:
                        # Unconditional set: apply to entire column
                        affected = len(new_df)
                        new_df[target_col] = new_value
                        summary = f"AI locally set '{target_col}' = '{new_value}' for all {affected} row(s)."
                    return new_df, summary
                else:
                    last_error = f"AI generated invalid set_value_intent: target_col='{target_col}' not found or not permitted."
                    break

            # Apply updates with column allowlist guard
            new_df = df.copy()
            affected = 0
            blocked_fields = set()

            def _apply_updates(update_list, source_df):
                """Apply a list of AI updates to new_df. Returns count of cells changed."""
                nonlocal affected
                row_map = {str(r.get('#', '')): idx for idx, r in source_df.iterrows()}
                for update in update_list:
                    row_serial = str(update.serial_number)
                    if row_serial and row_serial in row_map:
                        row_idx = row_map[row_serial]
                        update_dict = update.model_dump(exclude_unset=True, by_alias=True)
                        for k, v in update_dict.items():
                            if k == '#' or k == 'serial_number' or k == '::action': continue
                            if k in AI_ALLOWED_EDIT_COLS:
                                if new_df.at[row_idx, k] != v:
                                    new_df.at[row_idx, k] = v
                                    affected += 1
                            else:
                                blocked_fields.add(k)

            # Process first chunk updates
            if res_data.updates is not None:
                _apply_updates(res_data.updates, context_data_df)
                            
            # Process remaining chunks if dataset was truncated
            if _is_truncated and len(df) > len(context_data_df):
                import concurrent.futures
                
                def _process_chunk(chunk_df):
                    c_json = chunk_df.to_json(orient='records')
                    c_prompt = _build_prompt(c_json)
                    try:
                        c_comp = client.chat.completions.parse(
                            model=model,
                            messages=[
                                {"role": "system", "content": "You are a data patching engine."},
                                {"role": "user", "content": c_prompt}
                            ],
                            response_format=AISmartResponse,
                            timeout=60,
                            temperature=0.0
                        )
                        c_res = c_comp.choices[0].message.parsed
                        if c_res and c_res.updates:
                            return c_res.updates
                        return []
                    except Exception:
                        return []
                        
                chunk_size = len(context_data_df)
                chunks = [df.iloc[i:i+chunk_size] for i in range(chunk_size, len(df), chunk_size)]
                with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
                    futures = {executor.submit(_process_chunk, c): c for c in chunks}
                    for future in concurrent.futures.as_completed(futures):
                        c_df = futures[future]
                        c_updates = future.result()
                        if isinstance(c_updates, list):
                            _apply_updates(c_updates, c_df)

            summary = f"AI applied changes to {len(df)} row(s) ({affected} cell update(s))."
            if blocked_fields:
                summary += f" ⚠️ Blocked invalid field(s): {', '.join(sorted(blocked_fields))}."
            return new_df, summary

        except Exception as e:
            err_str = str(e)
            last_error = err_str[:200]
            log.warning("Smart Context attempt %d %s: %s", attempt + 1, type(e).__name__, err_str)
            # Rate limit or transient: back off and retry
            if "429" in err_str or "timeout" in err_str.lower() or "connection" in err_str.lower():
                wait = AI_RETRY_BASE_WAIT * (2 ** attempt)
                log.info("Retrying in %ds...", wait)
                time.sleep(wait)
                continue
            break  # Non-retryable error

    return None, f"AI command failed after {AI_RETRY_ATTEMPTS} attempt(s): {last_error}"
