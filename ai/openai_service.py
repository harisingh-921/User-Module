# user_masters/ai/openai_service.py
import json
import re
import base64
import streamlit as st
import pandas as pd
import fitz
import time
import hashlib
import io
import docx
from openai import OpenAI
from config.constants import USER_MASTER_COLS
from models.schemas import UserMasterResult, UserField

# ── Enterprise Safety Limits (edit here to tune) ─────────────────────────────
_MAX_FILE_SIZE_MB   = 20      # Hard reject uploads larger than this
_MAX_PDF_PAGES      = 60      # Cap PDF pages sent to AI
_MAX_AI_CONTEXT_KB  = 80      # Approx token guard: skip AI call if context > 80 KB
_AI_RETRY_ATTEMPTS  = 3       # Transient error retries for apply_ai_smart_context
_AI_RETRY_BASE_WAIT = 2       # Base seconds for exponential backoff
# Columns the AI is permitted to modify via apply_ai_smart_context
_AI_ALLOWED_EDIT_COLS = {
    'firstName', 'middleName', 'lastName', 'userName', 'email', 'mobile',
    'employeeId', 'departments', 'roles', 'units', 'designation',
    'isEnabled', 'password',
}
# ─────────────────────────────────────────────────────────────────────────────

USER_EXTRACTION_PROMPT = """
You are a High-Precision, Format-Agnostic User Data Extraction Engine.
Your task is to extract a structured user list from ANY document format — staff rosters, HR matrices, Word tables, PDFs, CSV exports, or any custom layout.

=== STEP 1: UNDERSTAND THE STRUCTURE ===
Before extracting, ANALYZE the document layout:
- Is this a STANDARD ROSTER? (one row per person with named columns like Name, Email, Department...)
- Is this a ROLE MATRIX? (columns represent roles, cells are Yes/No, and a separate column has the person name)
- Is this a FREE-FORMAT LIST? (names and details listed without a strict table structure)
- Could there be MERGED HEADERS or multi-row headers above the data?
DO NOT assume any fixed column order or name. Identify each column by its CONTENT and SEMANTIC MEANING.

=== STEP 2: FIELD IDENTIFICATION ===
Identify fields by MEANING, not by position or exact column header name:
- NAME field: Contains full person names including titles (e.g. "Dr. NG SWEET MAY", "Mr. RAMESH JOSHI"). Capture prefixes as part of the first name.
- EMAIL field: Contains @ symbols
- PHONE/MOBILE: Contains 010/011/012 patterns or dashes like 012-345-6789
- EMPLOYEE ID: Alphanumeric codes like CA00001, GB10007, EMP001
- DEPARTMENT: Hospital departments (Emergency Room, ICU, OPD, Surgery...)
- UNIT/FACILITY: Hospital branches or facility names (CAH Bukit Jalil, Ankura...)
- DESIGNATION/POSITION: Job titles (Doctor, Nurse, Consultant, Locum MO...)
- ROLES: Either a column with slash-separated text OR multiple Yes/No columns whose HEADERS are role names

=== STEP 3: EXTRACTION RULES ===
1. ROW INTEGRITY (MANDATORY): You MUST extract EVERY person in the data. Do not skip anyone. Do not summarize.
2. MULTI-USER SPLIT RULE (MANDATORY): A single row may contain data for MULTIPLE people, separated by a CONSISTENT DELIMITER. The delimiter can be ANY of: "|", ",", ";", "/", " & ", " and " — whichever is used consistently across that row.
   DETECTION: If the SAME delimiter appears in MULTIPLE columns of the same row (e.g. firstName has "Bijay | Maumita" AND employeeId has "1035708 | 192004"), that row contains multiple users. You MUST create SEPARATE JSON objects for each person.
   SPLITTING RULES:
   - Identify the delimiter being used (e.g. "|" or "," or "&").
   - User 1 takes the 1st part of EVERY delimited column.
   - User 2 takes the 2nd part of EVERY delimited column.
   - For columns with NO delimiter (e.g. middleName "Krishna" when others use "|"), apply the value ONLY to User 1. Leave it BLANK for User 2.
   - NEVER include the delimiter symbol in your final JSON fields — strip it out completely.
   - WORKED EXAMPLE with "|" delimiter (follow exactly):
       Source row: firstName="Bijay | Maumita", middleName="Krishna", lastName="Maity | Saha", employeeId="1035708 | 192004"
       firstName has "|" → 2 users. lastName has "|" → 2 users. employeeId has "|" → 2 users. middleName has NO "|" → belongs to User 1 only.
       Correct output:
         User 1: firstName="Bijay",   middleName="Krishna", lastName="Maity", employeeId="1035708"
         User 2: firstName="Maumita", middleName="",        lastName="Saha",  employeeId="192004"
   - WORKED EXAMPLE with "|" delimiter and MISSING lastName/employeeId for User 2 (follow exactly):
        Source row: firstName="Arindam | Riya", middleName="", lastName="Chakraborty", employeeId="10092"
        firstName has "|" → 2 users. middleName has NO "|" → blank for both. lastName has NO "|" → belongs ONLY to User 1. employeeId has NO "|" → belongs ONLY to User 1.
        Correct output:
          User 1: firstName="Arindam", middleName="", lastName="Chakraborty", employeeId="10092"
          User 2: firstName="Riya",    middleName="", lastName="",            employeeId=""
        WRONG output (never do this):
          User 2: firstName="Riya", lastName="Chakraborty"  ← copying no-pipe lastName to User 2 is FORBIDDEN.
          User 2: employeeId="10092" ← copying no-pipe employeeId to User 2 is FORBIDDEN.
   - WORKED EXAMPLE with 3-user split and NO middle names (follow exactly):
       Source row: firstName="Priyadarshini | Manisha | Anasuya", middleName="", lastName="Roy | Santra | Bajpaye"
       Correct output:
         User 1: firstName="Priyadarshini", lastName="Roy"
         User 2: firstName="Manisha",      lastName="Santra"
         User 3: firstName="Anasuya",     lastName="Bajpaye"
       WRONG output (never do this):
         User 1: firstName="Priyadarshini Manisha", lastName="Roy"  ← Merging names is FORBIDDEN.
   - WORKED EXAMPLE with "," delimiter:
       Source row: firstName="Anjali, Priya", lastName="Sen, Sharma", employeeId="101, 102"
       Correct output:
         User 1: firstName="Anjali", lastName="Sen",    employeeId="101"
         User 2: firstName="Priya",  lastName="Sharma", employeeId="102"
 3. FIELD MAPPING:
    - firstName, middleName, lastName: Extract logically from the source data.
    - STRICT COLUMN ALIGNMENT (CRITICAL): Do NOT shift data across columns. If the "Middle Name" column in Excel is blank for a specific user, the "middleName" field in JSON MUST be blank. NEVER move the second person's first name (from a pipe-split) into the first person's middleName field.
    - If the document has a "Full Name" column, split it logically into firstName and lastName.
    - employeeId: Extract the unique ID.
    - roles: This is CRITICAL. In this Excel format, the roles are often found in Column A as "floating headers" (e.g. 'Audit User - CESC - CAUTI') or are provided in the [ROLE SECTION:] tag at the start of the row. Combine all applicable roles found in Column A and [ROLE SECTION:] into this field, separated by "|".
    - departments / units / designation / email / mobile: Map correctly based on column headers.
4. UNIT NAME: Look at the COLUMN HEADERS for a global "Unit Name" (e.g. Unit Name | Anandpur) and apply it to every user.
5. NO LOGIC: Do not try to generate usernames or complex passwords. Just extract the raw name parts. Python will handle the rest.
6. NA LOGIC: If a value is "-", leave it blank.
7. ROLE SECTION TAGS (CRITICAL): Lines tagged with [ROLE SECTION: ...] are SECTION HEADERS, NOT people. NEVER create a user record from a [ROLE SECTION:] tag. Only use them to determine the `roles` value for the real data rows that follow. If a row has no valid name or employee ID, SKIP it entirely.
8. MIDDLE NAME RULE (CRITICAL): The `middleName` field is ONLY for a person's own middle name or initial (e.g. "K", "Kumar", "Rani"). It must NEVER contain another person's first name from a pipe-split. If firstName is "Arindam | Riya", you must create TWO separate users: User 1 has firstName="Arindam" and middleName="" (blank). User 2 has firstName="Riya" and middleName="" (blank). The text after the "|" is a second person, not a middle name.

OUTPUT FORMAT (JSON):
{{
  "document_name": "",
  "users": [
    {{
      "userName": "",
      "firstName": "",
      "middleName": "",
      "lastName": "",
      "employeeId": "",
      "departments": "",
      "roles": "",
      "units": "",
      "designation": "",
      "email": "",
      "mobile": "",
      "isEnabled": "Yes"
    }}
  ]
}}
"""

def _normalize_file(raw_df: pd.DataFrame) -> pd.DataFrame:
    """
    Converts any file format into a standard flat roster DataFrame.
    For Yes/No role-matrix files: detects the role columns, collapses them
    into a pipe-separated 'roles' column, and returns a clean DataFrame.
    For standard roster files: auto-detects headers and returns as-is.
    """
    # Stringify everything for analysis
    str_df = raw_df.astype(str).applymap(lambda x: x.strip())
    
    # ── Detect Yes/No matrix format ──────────────────────────────────────
    # Count Yes/No values per column
    yes_no_cols = []
    for col in str_df.columns:
        vals = str_df[col].str.lower()
        yn_count = vals.isin(['yes', 'no']).sum()
        if yn_count >= 2:  # At least 2 yes/no values = role column
            yes_no_cols.append(col)
    
    if len(yes_no_cols) >= 2:
        # This is a Yes/No matrix. Find the header row:
        # The header row is the last row BEFORE the first Yes/No data row.
        first_yn_row = None
        for idx, row in str_df.iterrows():
            if row.str.lower().isin(['yes', 'no']).any():
                first_yn_row = idx
                break
        
        if first_yn_row is not None and first_yn_row > 0:
            # The row just before the first Yes/No row is the role name row
            role_name_row_idx = first_yn_row - 1
            role_name_row = str_df.loc[role_name_row_idx]
            
            # Build a map: col_index → role_name (only for yes/no cols)
            role_col_map = {}
            for col in yes_no_cols:
                role_name = role_name_row[col].strip()
                if role_name and role_name.lower() not in ('nan', '-', ''):
                    role_col_map[col] = role_name
            
            # Find the name/user column (non-yes-no column with names)
            other_cols = [c for c in raw_df.columns if c not in yes_no_cols]
            
            # Data rows = rows from first_yn_row onwards
            data_df = raw_df.loc[first_yn_row:].copy().reset_index(drop=True)
            data_str = str_df.loc[first_yn_row:].reset_index(drop=True)
            
            # Build roles column per row
            def get_roles(row_idx):
                assigned = []
                for col, role_name in role_col_map.items():
                    val = data_str.loc[row_idx, col].strip().lower()
                    if val == 'yes':
                        assigned.append(role_name)
                return '|'.join(assigned)
            
            data_df['roles'] = [get_roles(i) for i in range(len(data_df))]
            
            # Drop individual role columns, keep others + new roles col
            data_df = data_df.drop(columns=list(role_col_map.keys()), errors='ignore')
            
            # Use role name row values as column headers for other cols
            new_cols = {}
            for col in other_cols:
                label = role_name_row[col].strip()
                if label and label.lower() not in ('nan', '-', ''):
                    new_cols[col] = label
                else:
                    new_cols[col] = str(col)
            data_df = data_df.rename(columns=new_cols)
            
            return data_df
    
    # ── Standard format: auto-detect header row ───────────────────────────
    # Find the first row that looks like a header (mostly text, not Yes/No)
    for idx, row in str_df.iterrows():
        vals = row.str.lower()
        if not vals.isin(['yes', 'no', 'nan', '-']).all():
            header = raw_df.loc[idx].astype(str).str.strip().tolist()
            data = raw_df.loc[idx+1:].copy().reset_index(drop=True)
            data.columns = header
            return data
    
    # Fallback: return as-is with auto column names
    raw_df.columns = [f"col_{i}" for i in range(len(raw_df.columns))]
    return raw_df


def _merge_duplicate_users(df: pd.DataFrame, pass_prefix: str = "Aone") -> pd.DataFrame:
    """
    After AI extraction, merge rows for the same user into one.
    Identifies duplicates by employeeId (preferred) or userName.
    Combines roles from all duplicate rows using '|' separator.
    """
    if df.empty:
        return df

    df = df.copy()

    # ── Drop phantom rows created from role section headers ──────────────────
    # A real user MUST have a name (firstName/lastName) or an employeeId.
    # If userName looks like a role string (e.g. "Audit User - MOS - ...") 
    # and has no employeeId, it's a misextracted header — remove it.
    ROLE_HEADER_KEYWORDS = ['audit user', 'incharge', 'audit incharge', 'med admin', 'ot manager', 
                             'nursing incharge', 'quality', 'infection control', 'micu', 'sicu', 'ccu',
                             'dialysis', 'radiology', 'emergency', 'lab', 'icn', 'icu', 'ot']
    def _is_phantom_row(row):
        emp = str(row.get('employeeId', '')).strip().lower()
        has_emp_id = emp and emp not in ('nan', 'none', '-', '')
        
        first = str(row.get('firstName', '')).strip().lower()
        last = str(row.get('lastName', '')).strip().lower()
        uname = str(row.get('userName', '')).strip().lower()
        
        has_first = first and first not in ('nan', 'none', '-', '')
        has_last = last and last not in ('nan', 'none', '-', '')
        has_uname = uname and uname not in ('nan', 'none', '-', '')
        
        # If absolutely no name fields and no employee ID, it's empty
        if not (has_first or has_last or has_uname) and not has_emp_id:
            return True
            
        if has_emp_id:
            return False  # Has an employee ID → definitely a real user
            
        # Check if name fields look like role strings
        for kw in ROLE_HEADER_KEYWORDS:
            if kw in uname and ' - ' in uname:  # Role section tags have " - " separators
                return True
            if kw in first and not last:  # firstName = role keyword, no lastName
                return True
        return False

    df = df[~df.apply(_is_phantom_row, axis=1)].reset_index(drop=True)

    # Determine the key to group by
    def get_master_key(row):
        eid = str(row.get('employeeId', '')).strip().lower()
        if eid and eid not in ('nan', 'none', '-', ''):
            return f"id_{eid}"
        
        email = str(row.get('email', '')).strip().lower()
        if email and email not in ('nan', 'none', '-', ''):
            return f"email_{email}"
            
        mobile = str(row.get('mobile', '')).strip().lower()
        if mobile and mobile not in ('nan', 'none', '-', ''):
            return f"mobile_{mobile}"
            
        uname = str(row.get('userName', '')).strip().lower()
        if uname and uname not in ('nan', 'none', '-', ''):
            return f"user_{uname}"
            
        first = str(row.get('firstName', '')).strip().lower()
        last = str(row.get('lastName', '')).strip().lower()
        if first and first not in ('nan', 'none', '-', '') and last and last not in ('nan', 'none', '-', ''):
            f_clean = re.sub(r'[^a-z]', '', first)
            l_clean = re.sub(r'[^a-z]', '', last)
            if f_clean or l_clean:
                return f"name_{f_clean}_{l_clean}"
            
        return f"unkeyed_{row.name}"

    df['_group_key'] = df.apply(get_master_key, axis=1)

    def merge_group(group):
        # Prefer the "cleanest" row: sort by firstName length ascending so a short,
        # correctly-split name (e.g. "Arindam") wins over a merged one (e.g. "Arindam Riya").
        # This does NOT affect single-occurrence users at all — sorting a 1-row group is a no-op.
        if len(group) > 1 and 'firstName' in group.columns:
            group = group.copy()
            group['_fn_len'] = group['firstName'].astype(str).str.replace('|', '', regex=False).str.strip().str.len()
            group = group.sort_values('_fn_len').drop(columns=['_fn_len'])
        merged = group.iloc[0].copy()
        # Merge all role strings across duplicates
        if 'roles' in group.columns:
            all_roles = []
            for r in group['roles'].dropna():
                for part in str(r).split('|'):
                    part = part.strip()
                    if part and part.lower() not in ('nan', '-', ''):
                        if part not in all_roles:
                            all_roles.append(part)
            merged['roles'] = '|'.join(all_roles)
        # Fill blank fields from other rows in the group
        for col in group.columns:
            if col in ('roles', '_group_key'):
                continue
            if str(merged.get(col, '')).strip().lower() in ('', 'nan', 'none', '-'):
                for val in group[col].dropna():
                    if str(val).strip().lower() not in ('', 'nan', 'none', '-'):
                        merged[col] = val
                        break
        return merged

    # Use an explicit loop instead of groupby().apply() — in pandas 2.x,
    # apply() with a function returning a Series can silently drop single-occurrence rows.
    result_rows = []
    for _, group in df.groupby('_group_key', sort=False):
        result_rows.append(merge_group(group))

    merged_df = pd.DataFrame(result_rows).reset_index(drop=True)
    merged_df = merged_df.drop(columns=['_group_key'], errors='ignore')
    
    # Programmatic userName construction: join first+middle+last and keep ONLY letters
    def construct_username(row):
        parts = []
        fn = str(row.get('firstName', '')).strip()
        mn = str(row.get('middleName', '')).strip()
        ln = str(row.get('lastName', '')).strip()

        # Safety 1: Stop at pipes (Primary split logic)
        if '|' in fn: fn = fn.split('|')[0].strip()
        if '|' in mn: mn = mn.split('|')[0].strip()
        if '|' in ln: ln = ln.split('|')[0].strip()

        # Safety 2: Anti-Merge (If AI returned "Priyodarshini Manisha" without a pipe)
        # We take only the first word for the first name to prevent merging users.
        if ' ' in fn:
            fn = fn.split(' ')[0].strip()

        for val in [fn, mn, ln]:
            if val and val.lower() not in ('nan', 'none', '-'):
                parts.append(val)
        
        full = "".join(parts)
        # Keep only letters, lowercase
        clean = re.sub(r'[^a-zA-Z]', '', full).lower()
        return clean if clean else row.get('userName', 'user')

    # Programmatic password generation
    def construct_password(row):
        emp_id = str(row.get('employeeId', '')).strip()
        if not emp_id or emp_id.lower() in ('nan', 'none', '-'):
            return "" # Keep password blank if employeeId is missing
        return f"{pass_prefix}@{emp_id}"

    merged_df['userName'] = merged_df.apply(construct_username, axis=1)
    merged_df['password'] = merged_df.apply(construct_password, axis=1)

    # ── DOUBLE MERGE PASS ────────────────────────────────────────────────────
    # Now that usernames are clean (no more 'priyodarshinimanisha'), 
    # we run the merger AGAIN. This ensures that a previously 'merged' name 
    # and a 'clean' name are now recognized as the same person.
    final_rows = []
    merged_df['_final_group_key'] = merged_df['userName']
    for _, group in merged_df.groupby('_final_group_key', sort=False):
        if len(group) == 1:
            final_rows.append(group.iloc[0])
        else:
            # Re-run the role merge logic
            m = group.iloc[0].copy()
            all_roles = []
            for r in group['roles'].dropna():
                for part in str(r).split('|'):
                    part = part.strip()
                    if part and part.lower() not in ('nan', '-', '') and part not in all_roles:
                        all_roles.append(part)
            m['roles'] = '|'.join(all_roles)
            final_rows.append(m)
    
    merged_df = pd.DataFrame(final_rows).reset_index(drop=True)
    merged_df = merged_df.drop(columns=['_final_group_key'], errors='ignore')
    # ─────────────────────────────────────────────────────────────────────────
        
    return merged_df


def find_matching_excel_roles(user_dict, excel_rows_data):
    emp_id = str(user_dict.get('employeeId', '')).strip().lower()
    email = str(user_dict.get('email', '')).strip().lower()
    uname = str(user_dict.get('userName', '')).strip().lower()
    
    # Priority 1: Match by Employee ID
    if emp_id and emp_id not in ('nan', 'none', '-', ''):
        for row_info in excel_rows_data:
            if emp_id in row_info['raw_values']:
                return row_info['roles']
                
    # Priority 2: Match by Email
    if email and email not in ('nan', 'none', '-', ''):
        for row_info in excel_rows_data:
            if email in row_info['raw_values']:
                return row_info['roles']
                
    # Priority 3: Match by Username
    if uname and uname not in ('nan', 'none', '-', ''):
        for row_info in excel_rows_data:
            if uname in row_info['raw_values'] or any(uname in str(val) for val in row_info['raw_values']):
                return row_info['roles']
                
    # Fallback: return whatever roles the LLM extracted
    return user_dict.get('roles', '')

def get_all_api_keys(primary_key=None):
    keys = []
    if primary_key:
        p_str = str(primary_key).strip()
        if p_str and p_str not in keys:
            keys.append(p_str)
            
    # Load from st.secrets
    try:
        # Check standard names
        for k in ["OPENAI_API_KEY", "OPENAI_API_KEY_2", "OPENAI_API_KEY_3", "GEMINI_API_KEY", "GEMINI_API_KEY_2", "GEMINI_API_KEY_3"]:
            val = str(st.secrets.get(k, "")).strip()
            if val and val not in keys:
                keys.append(val)
                
        # Also grab any key matching dynamic pattern
        for k in st.secrets.keys():
            val = str(st.secrets.get(k, "")).strip()
            if (val.startswith("sk-") or val.startswith("AIzaSy")) and val not in keys:
                keys.append(val)
    except Exception:
        pass
    return keys


def local_extract_users(file_bytes, filename, pass_prefix="Aone"):
    """
    LOCAL extraction engine — NO AI, NO API calls.
    Reads Excel/CSV, auto-detects headers, maps columns to our schema using
    fuzzy matching against SEMANTIC_MAPPINGS, and returns a clean DataFrame.
    Works even when all API keys are exhausted.
    """
    from config.constants import USER_MASTER_COLS, SEMANTIC_MAPPINGS
    
    ext = filename.lower()
    try:
        if ext.endswith('.csv'):
            all_sheets = {'Sheet1': pd.read_csv(io.BytesIO(file_bytes), header=None)}
        elif ext.endswith(('.xlsx', '.xls')):
            all_sheets = pd.read_excel(io.BytesIO(file_bytes), header=None, sheet_name=None)
        else:
            st.warning(f"Local extraction only supports Excel/CSV. '{filename}' skipped.")
            return pd.DataFrame()
    except Exception as e:
        st.error(f"Error reading {filename}: {e}")
        return pd.DataFrame()
    
    all_users = []
    
    for sheet_name, raw_df in all_sheets.items():
        raw_df = raw_df.dropna(how='all')
        mask = raw_df.astype(str).apply(lambda x: x.str.contains(r'[a-zA-Z0-9]', na=False)).any(axis=1)
        raw_df = raw_df[mask].reset_index(drop=True)
        if raw_df.empty:
            continue
            
        # Try to find a global unit name in the sheet
        global_unit = ""
        for r_idx, row in raw_df.iterrows():
            row_vals = [str(x).strip() for x in row.values]
            for c_idx, val in enumerate(row_vals):
                if val.lower() == 'unit name' and c_idx + 1 < len(row_vals):
                    candidate = row_vals[c_idx + 1]
                    if candidate and candidate.lower() not in ('nan', 'none', '-'):
                        global_unit = candidate
                        break
            if global_unit:
                break
        
        # --- Improved Auto-detect header row ---
        str_df = raw_df.astype(str).map(lambda x: str(x).strip())
        header_row_idx = 0
        max_matches = 0
        for idx, row in str_df.iterrows():
            vals = row.str.lower().tolist()
            # If row contains common header keywords, it's likely the header
            header_keywords = ['name', 'email', 'employee', 'id', 'mobile', 'phone', 
                             'department', 'unit', 'role', 'designation', 'staff',
                             'first', 'last', 'username', 'password']
            matches = sum(1 for v in vals if any(kw in str(v).lower() for kw in header_keywords))
            if matches > max_matches:
                max_matches = matches
                header_row_idx = idx
            if matches >= 5:
                break
        
        # --- Check if the row immediately following the header row is a sub-header row ---
        is_sub_header = False
        headers_temp = [str(h).strip() for h in raw_df.iloc[header_row_idx].values]
        headers_lower_temp = {str(h).strip(): str(h).lower().strip() for h in raw_df.iloc[header_row_idx].values}
        
        # Build temp col_mapping to check if next row has actual data in name/email columns
        col_mapping_temp = {}
        for target_field in USER_MASTER_COLS:
            if target_field == 'roles': continue
            tf_lower = target_field.lower()
            for src_col, src_lower in headers_lower_temp.items():
                if src_col in col_mapping_temp: continue
                if src_lower == tf_lower or src_lower.replace(' ', '') == tf_lower.lower():
                    col_mapping_temp[src_col] = target_field
                    break
            else:
                if target_field in SEMANTIC_MAPPINGS:
                    for alias in SEMANTIC_MAPPINGS[target_field]:
                        for src_col, src_lower in headers_lower_temp.items():
                            if src_col in col_mapping_temp: continue
                            if alias in src_lower or src_lower in alias:
                                col_mapping_temp[src_col] = target_field
                                break
                        if any(v == target_field for v in col_mapping_temp.values()):
                            break
                            
                if not any(v == target_field for v in col_mapping_temp.values()):
                    broad_keywords = {
                        'departments': ['department', 'dept'],
                        'units': ['unit', 'ward', 'division'],
                        'designation': ['designation', 'position', 'title', 'rank', 'category'],
                        'userName': ['user name', 'username'],
                        'employeeId': ['employee id', 'emp id', 'staff id', 'emp no', 'employee no', 'id no'],
                        'email': ['email', 'e-mail', 'mail'],
                        'mobile': ['mobile', 'phone', 'contact', 'cell', 'telephone'],
                    }
                    if target_field in broad_keywords:
                        for kw in broad_keywords[target_field]:
                            for src_col, src_lower in headers_lower_temp.items():
                                if src_col in col_mapping_temp: continue
                                if kw in src_lower:
                                    col_mapping_temp[src_col] = target_field
                                    break
                            if any(v == target_field for v in col_mapping_temp.values()):
                                break
                                
                if target_field == 'firstName' and 'firstName' not in col_mapping_temp.values():
                    for src_col, src_lower in headers_lower_temp.items():
                        if src_col in col_mapping_temp: continue
                        if src_lower in ('name', 'full name', 'fullname', 'staff name', 'employee name'):
                            col_mapping_temp[src_col] = '_fullName'
                            break

        if header_row_idx + 1 < len(raw_df):
            next_row = raw_df.iloc[header_row_idx + 1]
            name_email_empty = True
            for src_col, target_field in col_mapping_temp.items():
                if target_field in ['firstName', 'lastName', '_fullName', 'employeeId', 'email']:
                    col_index = raw_df.iloc[header_row_idx].tolist().index(src_col)
                    val = str(next_row.iloc[col_index]).strip().lower() if col_index < len(next_row) else ""
                    if val and val not in ('nan', 'none', '-'):
                        name_email_empty = False
                        break
            text_cells = sum(1 for v in next_row.values if str(v).strip() and str(v).strip().lower() not in ('nan', 'none', '-'))
            if name_email_empty and text_cells >= 2:
                is_sub_header = True

        if is_sub_header:
            headers = []
            parent_headers = raw_df.iloc[header_row_idx].tolist()
            # Forward-fill parent headers to handle merged cells!
            filled_parents = []
            last_parent = ""
            for p in parent_headers:
                p_str = str(p).strip()
                if p_str and p_str.lower() not in ('nan', 'none', '-'):
                    last_parent = p_str
                filled_parents.append(last_parent)

            for c_idx in range(len(raw_df.columns)):
                parent_h = filled_parents[c_idx]
                child_h = str(raw_df.iloc[header_row_idx + 1].iloc[c_idx]).strip()
                
                parent_clean = "" if parent_h.lower() in ('nan', 'none', '-') else parent_h
                child_clean = "" if child_h.lower() in ('nan', 'none', '-') else child_h
                
                if parent_clean and child_clean:
                    if parent_clean.lower() == child_clean.lower():
                        headers.append(child_clean)
                    else:
                        headers.append(f"{parent_clean}|{child_clean}")
                elif child_clean:
                    headers.append(child_clean)
                elif parent_clean:
                    headers.append(parent_clean)
                else:
                    headers.append(f"col_{c_idx}")
            data_df = raw_df.iloc[header_row_idx + 2:].copy().reset_index(drop=True)
        else:
            headers = [str(h).strip() for h in raw_df.iloc[header_row_idx].values]
            data_df = raw_df.iloc[header_row_idx + 1:].copy().reset_index(drop=True)

        # Ensure unique headers by adding suffixes to duplicates to prevent pandas Series mapping issues
        unique_headers = []
        header_counts = {}
        for h in headers:
            if h in header_counts:
                header_counts[h] += 1
                unique_headers.append(f"{h}_{header_counts[h]}")
            else:
                header_counts[h] = 0
                unique_headers.append(h)
        headers = unique_headers
        data_df.columns = headers
        
        # Remove rows that are all empty/nan
        data_df = data_df.dropna(how='all').reset_index(drop=True)
        data_str = data_df.astype(str).map(lambda x: str(x).strip())
        has_content = data_str.apply(lambda row: row.str.lower().replace('nan', '').replace('none', '').replace('-', '').str.strip().ne('').any(), axis=1)
        data_df = data_df[has_content].reset_index(drop=True)
        
        if data_df.empty:
            continue
        
        # --- Map columns to our schema ---
        col_mapping = {}  # source_col -> target_field
        headers_lower = {h: h.lower().strip() for h in headers}
        
        for target_field in USER_MASTER_COLS:
            # Roles is handled specially if it's a running role column or tick-marked columns
            if target_field == 'roles':
                continue
            if any(v == target_field for v in col_mapping.values()):
                continue  # Already mapped
            tf_lower = target_field.lower()
            # Direct match
            for src_col, src_lower in headers_lower.items():
                if src_col in col_mapping:
                    continue
                if src_lower == tf_lower or src_lower.replace(' ', '') == tf_lower.lower():
                    col_mapping[src_col] = target_field
                    break
            else:
                # Semantic/fuzzy match using SEMANTIC_MAPPINGS
                if target_field in SEMANTIC_MAPPINGS:
                    for alias in SEMANTIC_MAPPINGS[target_field]:
                        for src_col, src_lower in headers_lower.items():
                            if src_col in col_mapping:
                                continue
                            if alias in src_lower or src_lower in alias:
                                col_mapping[src_col] = target_field
                                break
                        if any(v == target_field for v in col_mapping.values()):
                            break
                
                # Broad keyword match for common fields
                if not any(v == target_field for v in col_mapping.values()):
                    broad_keywords = {
                        'departments': ['department', 'dept', 'location', 'branch', 'facility', 'site'],
                        'units': ['unit', 'ward', 'section', 'division'],
                        'designation': ['designation', 'position', 'title', 'rank', 'category'],
                        'userName': ['user name', 'username', 'login', 'user id', 'userid'],
                        'employeeId': ['employee id', 'emp id', 'staff id', 'emp no', 'employee no', 'id no'],
                        'email': ['email', 'e-mail', 'mail'],
                        'mobile': ['mobile', 'phone', 'contact', 'cell', 'telephone'],
                    }
                    if target_field in broad_keywords:
                        for kw in broad_keywords[target_field]:
                            for src_col, src_lower in headers_lower.items():
                                if src_col in col_mapping:
                                    continue
                                if kw in src_lower:
                                    col_mapping[src_col] = target_field
                                    break
                            if any(v == target_field for v in col_mapping.values()):
                                break
                            
                # Handle "Full Name" / "Name" -> firstName + lastName split
                if target_field == 'firstName' and 'firstName' not in col_mapping.values():
                    for src_col, src_lower in headers_lower.items():
                        if src_col in col_mapping:
                            continue
                        if src_lower in ('name', 'full name', 'fullname', 'staff name', 'employee name'):
                            col_mapping[src_col] = '_fullName'
                            break
        
        # Check for running/floating roles column in headers (must not be a tick-marked column)
        TICK_VALUES = {'yes', 'y', 'x', '1', 'true', 'v', '\u221a', '\u2713', '\u2714', '\u2611'}  # includes √ ✓ ✔ ☑
        roles_col_name = None
        for h in headers:
            if any(kw in h.lower() for kw in ['role', 'audit user', 'assigned role', 'user role', 'incharge', 'admin', 'validation']):
                col_vals = data_df[h].dropna().astype(str).str.strip()
                has_ticks = col_vals.apply(lambda v: v.lower() in TICK_VALUES or v in TICK_VALUES).any()
                if not has_ticks:
                    roles_col_name = h
                    break

        # --- Detect tick-marked role columns ---
        role_cols = {}
        role_keywords = ['role', 'audit', 'incharge', 'admin', 'user', 'manager', 
                        'operator', 'reporter', 'viewer', 'approver', 'officer', 'staff', 'analyst', 'advisor', 'preventionist']
        for src_col in headers:
            src_lower = str(src_col).lower()
            is_role_header = any(kw in src_lower for kw in role_keywords)
            if is_role_header and src_col not in col_mapping and src_col != roles_col_name:
                col_vals = data_df[src_col].dropna().astype(str).str.strip()
                has_ticks = col_vals.apply(lambda v: v.lower() in TICK_VALUES or v in TICK_VALUES).any()
                if has_ticks:
                    role_cols[src_col] = src_col
        
        # --- Build user records ---
        last_role_val = ""
        for _, row in data_df.iterrows():
            user = {col: '' for col in USER_MASTER_COLS}
            
            # Global Unit Name fallback
            if global_unit:
                user['units'] = global_unit
                
            # Running role mapping
            if roles_col_name:
                rv = str(row.get(roles_col_name, '')).strip()
                if rv and rv.lower() not in ('nan', 'none', '-'):
                    last_role_val = rv
            
            if last_role_val:
                user['roles'] = last_role_val
            
            for src_col, target_field in col_mapping.items():
                val = str(row.get(src_col, '')).strip()
                if val.lower() in ('nan', 'none', '-'):
                    val = ''
                
                if target_field == '_fullName':
                    # Split "Full Name" into firstName + lastName
                    parts = val.split()
                    if len(parts) >= 2:
                        user['firstName'] = parts[0]
                        user['lastName'] = ' '.join(parts[1:])
                    elif len(parts) == 1:
                        user['firstName'] = parts[0]
                else:
                    user[target_field] = val
            
            # Collect tick-marked roles
            if role_cols:
                assigned_roles = []
                for rc_col, rc_name in role_cols.items():
                    rv = str(row.get(rc_col, '')).strip()
                    if rv.lower() in TICK_VALUES or rv in TICK_VALUES:
                        clean_rc_name = rc_name.split('|')[-1] if '|' in rc_name else rc_name
                        parent_pre = rc_name.split('|')[0] if '|' in rc_name else ""
                        if parent_pre and parent_pre.lower() not in clean_rc_name.lower():
                            assigned_roles.append(f"{parent_pre}|{clean_rc_name}")
                        else:
                            assigned_roles.append(clean_rc_name)
                if assigned_roles:
                    if user['roles']:
                        existing = user['roles'].split('|')
                        for r in assigned_roles:
                            if r not in existing:
                                existing.append(r)
                        user['roles'] = '|'.join(existing)
                    else:
                        user['roles'] = '|'.join(assigned_roles)
            
            # --- SPLIT MULTI-USER ROWS (PIPE SEPARATED DELIMITER) ---
            identity_fields = ['firstName', 'middleName', 'lastName', 'userName', 'employeeId', 'email', 'mobile']
            max_parts = 1
            for f in identity_fields:
                val = user.get(f, '')
                if '|' in val:
                    parts = [p.strip() for p in val.split('|')]
                    max_parts = max(max_parts, len(parts))
            
            if max_parts > 1:
                for i in range(max_parts):
                    sub_user = user.copy()
                    for f in USER_MASTER_COLS:
                        if f in identity_fields:
                            val = user.get(f, '')
                            parts = [p.strip() for p in val.split('|')] if '|' in val else [val]
                            if i < len(parts):
                                sub_user[f] = parts[i]
                            else:
                                sub_user[f] = ''
                    
                    has_name = (sub_user.get('firstName', '').strip() or 
                               sub_user.get('lastName', '').strip() or 
                               sub_user.get('employeeId', '').strip() or
                               sub_user.get('userName', '').strip())
                    if has_name:
                        all_users.append(sub_user)
            else:
                has_name = (user.get('firstName', '').strip() or 
                           user.get('lastName', '').strip() or 
                           user.get('employeeId', '').strip() or
                           user.get('userName', '').strip())
                if has_name:
                    all_users.append(user)
    
    if not all_users:
        return pd.DataFrame()
    
    raw_df = pd.DataFrame(all_users)
    result_df = _merge_duplicate_users(raw_df, pass_prefix=pass_prefix)
    return result_df


def get_openai_client(api_key):
    if not api_key:
        return None
    api_key_str = str(api_key).strip()
    if api_key_str.startswith("AIzaSy"):
        # This is a Gemini API Key! Use Gemini's OpenAI-compatible endpoint.
        return OpenAI(
            api_key=api_key_str,
            base_url="https://generativelanguage.googleapis.com/v1beta/openai/"
        )
    return OpenAI(api_key=api_key_str)

def openai_extract_users(file_bytes, filename, api_key, intent="", pass_prefix="Aone"):
    """Universal AI User Extraction Engine with Chunking & Failover."""
    client = get_openai_client(api_key)
    if not client: return None

    # ── File-size guard ───────────────────────────────────────────────────────
    size_mb = len(file_bytes) / 1_048_576
    if size_mb > _MAX_FILE_SIZE_MB:
        st.error(f"❌ **{filename}** is {size_mb:.1f} MB — exceeds the {_MAX_FILE_SIZE_MB} MB limit. "
                 f"Please split the file and re-upload.")
        return pd.DataFrame()
    # ─────────────────────────────────────────────────────────────────────────

    try:
        dynamic_prompt = USER_EXTRACTION_PROMPT.format(pass_prefix=pass_prefix)
        ext = filename.lower()
        
        if ext.endswith(('.xlsx', '.xls', '.csv')):
            if ext.endswith('.csv'):
                raw_df = pd.read_csv(io.BytesIO(file_bytes), header=None)
                sheet_dfs = {'Sheet1': raw_df}
            else:
                # Read ALL sheets — sheet_name=None returns {sheet_name: DataFrame}
                all_sheets = pd.read_excel(io.BytesIO(file_bytes), header=None, sheet_name=None)
                sheet_dfs = all_sheets
                st.toast(f"📋 Found {len(sheet_dfs)} sheet(s): {', '.join(sheet_dfs.keys())}")
            
            # We collect all chunks across all sheets, each with its own correct header context
            all_chunks_to_process = []
            
            global_excel_rows_data = []
            has_tick_role_columns = False
            
            for sheet_name, raw_df in sheet_dfs.items():
                raw_df = raw_df.dropna(how='all')
                mask = raw_df.astype(str).apply(lambda x: x.str.contains(r'[a-zA-Z0-9]', na=False)).any(axis=1)
                raw_df = raw_df[mask].reset_index(drop=True)
                if raw_df.empty: continue
                
                str_df = raw_df.astype(str).map(lambda x: str(x).strip())
                
                # Detect header row for THIS sheet
                first_data_row = -1
                for i, row in str_df.iterrows():
                    row_vals = row.str.lower()
                    if row_vals.isin(['yes', 'no']).any():
                        first_data_row = i
                        break
                    # Check for common header keywords
                    if any(kw in str(v).lower() for kw in ['name', 'email', 'employee', 'id', 'mobile', 'phone', 'department', 'unit', 'role'] for v in row.values if v):
                        first_data_row = i + 1
                        break
                
                if first_data_row == -1:
                    first_data_row = 1 if len(raw_df) > 1 else 0
                    
                header_rows_df = raw_df.iloc[:first_data_row]
                data_rows_df = raw_df.iloc[first_data_row:]
                
                def row_to_str(row):
                    return " ; ".join([str(v).strip().replace(';', ',') if str(v).strip().lower() not in ('nan','none','') else '-' for v in row.values])
                
                sheet_header_context = f"COLUMN HEADERS (from '{sheet_name}'):\n" + "\n".join(row_to_str(r) for _, r in header_rows_df.iterrows())
                
                # --- IDENTIFY TICK-MARKED ROLE COLUMNS IN PYTHON ---
                header_row_idx = first_data_row - 1 if first_data_row > 0 else 0
                headers = [str(h).strip() for h in raw_df.iloc[header_row_idx].values]
                
                role_cols = {}
                for col_idx, header in enumerate(headers):
                    header_lower = header.lower()
                    is_role_header = any(kw in header_lower for kw in ['role', 'audit', 'incharge', 'admin', 'user', 'manager', 'operator', 'reporter', 'viewer', 'approver', 'officer', 'staff'])
                    col_values = data_rows_df.iloc[:, col_idx].dropna().astype(str).str.strip().str.lower()
                    has_ticks = col_values.isin(['✓', '✔', 'yes', 'y', 'x', '1', 'true', 'v']).any()
                    if is_role_header and has_ticks:
                        role_cols[col_idx] = header
                        has_tick_role_columns = True
                
                # --- EXTRACT ROLE TICKS PER ROW ---
                for idx, row in data_rows_df.iterrows():
                    row_roles = []
                    for col_idx, role_name in role_cols.items():
                        val = str(row.iloc[col_idx]).strip().lower() if col_idx < len(row) else ""
                        if val in ('✓', '✔', 'yes', 'y', 'x', '1', 'true', 'v'):
                            row_roles.append(role_name)
                    
                    raw_vals = []
                    for v in row.values:
                        if pd.notna(v):
                            v_str = str(v).strip().lower()
                            if v_str not in ('nan', 'none', '', '-'):
                                raw_vals.append(v_str)
                    
                    global_excel_rows_data.append({
                        'roles': '|'.join(row_roles),
                        'raw_values': raw_vals
                    })
                
                sheet_lines = []
                last_col_a_value = ""
                for _, row in data_rows_df.iterrows():
                    line = row_to_str(row)
                    if not line.strip(): continue
                    col_a_val = str(row.iloc[0]).strip() if len(row) > 0 else ""
                    if col_a_val.lower() not in ('nan', 'none', '', '-'):
                        last_col_a_value = col_a_val
                    tag = f"[SHEET: {sheet_name}]"
                    if last_col_a_value:
                        tag += f" [ROLE SECTION: {last_col_a_value}]"
                    sheet_lines.append(f"{tag} {line}")
                
                # Create chunks for this specific sheet
                # Small chunks (10 rows) = higher precision, fewer skipped users
                chunk_size = 40
                for i in range(0, len(sheet_lines), chunk_size):
                    chunk = sheet_lines[i:i + chunk_size]
                    all_chunks_to_process.append({
                        'text': sheet_header_context + "\n\nDATA:\n" + "\n".join(chunk),
                        'sheet': sheet_name,
                        'line_start': i
                    })

            if not all_chunks_to_process:
                st.error("⚠️ No data found in any of the Excel sheets.")
                return pd.DataFrame()
            
            from concurrent.futures import ThreadPoolExecutor, as_completed
            progress_bar = st.progress(0)
            status_text = st.empty()
            status_text.info(f"🚀 Started extraction of {len(all_chunks_to_process)} batches...")
            
            all_users = []
            
            def process_chunk(chunk_data, batch_idx):
                chunk_text = chunk_data['text']
                all_keys = get_all_api_keys(api_key)
                
                for k_idx, current_key in enumerate(all_keys):
                    current_client = get_openai_client(current_key)
                    if not current_client:
                        continue
                    
                    if current_key.startswith("AIzaSy"):
                        models_to_try = ["gemini-2.5-flash"]
                    else:
                        models_to_try = ["gpt-4o-mini", "gpt-4o"]
                        
                    for model in models_to_try:
                        try:
                            completion = current_client.chat.completions.parse(
                                model=model,
                                messages=[
                                    {"role": "system", "content": dynamic_prompt},
                                    {"role": "user", "content": f"USER INTENT: {intent}\n\nINPUT DATA TABLE:\n{chunk_text}"}
                                ],
                                response_format=UserMasterResult,
                                timeout=120,
                                temperature=0.0
                            )
                            result = [u.__dict__ for u in completion.choices[0].message.parsed.users]
                            if not result and (model == "gpt-4o-mini" or model == "gemini-2.5-flash"):
                                time.sleep(1)
                                completion2 = current_client.chat.completions.parse(
                                    model=model,
                                    messages=[
                                        {"role": "system", "content": dynamic_prompt},
                                        {"role": "user", "content": f"USER INTENT: {intent}\n\nINPUT DATA TABLE:\n{chunk_text}"}
                                    ],
                                    response_format=UserMasterResult,
                                    timeout=120,
                                    temperature=0.0
                                )
                                result = [u.__dict__ for u in completion2.choices[0].message.parsed.users]
                            return result
                        except Exception as e:
                            err_str = str(e).lower()
                            if "quota" in err_str or "rate limit" in err_str or "429" in err_str:
                                print(f"[Batch {batch_idx}] API Key #{k_idx+1} hit rate limit/quota. Trying next key...")
                                break # break the model loop to try the next key
                            else:
                                if model == models_to_try[0]:
                                    time.sleep(1)
                                    continue # try next model
                                print(f"Batch {batch_idx} AI Error: {e}")
                return []

            # Gemini Free Tier has a strict 15 RPM / concurrency limit.
            # Running sequentially (max_workers=1) ensures we never hit the rate limit!
            max_workers = 1 if api_key_str.startswith("AIzaSy") else 5
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                future_to_idx = {executor.submit(process_chunk, chunk, i): i for i, chunk in enumerate(all_chunks_to_process)}
                completed = 0
                # Store results by index to preserve source order
                ordered_results = {}
                for future in as_completed(future_to_idx):
                    idx = future_to_idx[future]
                    ordered_results[idx] = future.result()
                    completed += 1
                    progress_bar.progress(completed / len(all_chunks_to_process))
                    captured = sum(len(v) for v in ordered_results.values())
                    status_text.info(f"⚡ Extraction: {completed}/{len(all_chunks_to_process)} batches done ({captured} users captured)")
            
            # Flatten in chunk order to preserve original row sequence
            all_users = []
            for i in sorted(ordered_results.keys()):
                all_users.extend(ordered_results[i])
            
            status_text.empty()
            progress_bar.empty()
            
            # Override roles with exact Excel ticks if available
            if has_tick_role_columns and global_excel_rows_data:
                for user in all_users:
                    matched_roles = find_matching_excel_roles(user, global_excel_rows_data)
                    user['roles'] = matched_roles
            
            raw_df = pd.DataFrame(all_users)
            result_df = _merge_duplicate_users(raw_df, pass_prefix=pass_prefix)
            st.toast(f"✅ {len(raw_df)} raw rows → {len(result_df)} unique users after merge.")
            return result_df

        elif ext.endswith('.pdf'):
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            total_pages = len(doc)
            if total_pages > _MAX_PDF_PAGES:
                st.warning(f"⚠️ **{filename}** has {total_pages} pages — processing first {_MAX_PDF_PAGES} only.")
            all_lines = []
            for page in doc.pages(0, min(total_pages, _MAX_PDF_PAGES)):
                text = page.get_text("text")
                lines = [line.strip() for line in text.split("\n") if line.strip()]
                all_lines.extend(lines)
            st.toast(f"📄 Read {min(total_pages, _MAX_PDF_PAGES)}/{total_pages} PDF pages.")
            header_context = "SOURCE: PDF Document\n"
            
        elif ext.endswith(('.docx', '.doc')):
            doc = docx.Document(io.BytesIO(file_bytes))
            all_lines = []
            for para in doc.paragraphs:
                if para.text.strip():
                    all_lines.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    if any(row_data):
                        all_lines.append(" | ".join(row_data))
            st.toast(f"📄 Read Word document.")
            header_context = "SOURCE: Word Document\n"
            
        else:
            st.error(f"Unsupported file format: {ext}")
            return None

        # Fallback for PDF/Word
        chunk_size = 30
        all_users = []
        chunks = []
        for i in range(0, len(all_lines), chunk_size):
            chunk = all_lines[i:i + chunk_size]
            chunks.append(header_context + "\n\nDATA:\n" + "\n".join(chunk))
        
        if not chunks: return pd.DataFrame()
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        def process_simple_chunk(chunk_text, batch_idx):
            for model in ["gpt-4o", "gpt-4o-mini"]:
                try:
                    completion = client.chat.completions.parse(
                        model=model,
                        messages=[
                            {"role": "system", "content": dynamic_prompt},
                            {"role": "user", "content": f"USER INTENT: {intent}\n\nINPUT DATA TABLE:\n{chunk_text}"}
                        ],
                        response_format=UserMasterResult,
                        timeout=120,
                        temperature=0.0
                    )
                    return [u.__dict__ for u in completion.choices[0].message.parsed.users]
                except Exception as e:
                    # Log error but continue to next model — never silently swallow
                    print(f"[PDF/Word Batch {batch_idx}] model={model} error: {type(e).__name__}: {e}")
                    continue
            return []

        with ThreadPoolExecutor(max_workers=5) as executor:
            future_to_idx = {executor.submit(process_simple_chunk, text, i): i for i, text in enumerate(chunks)}
            completed = 0
            for future in as_completed(future_to_idx):
                all_users.extend(future.result())
                completed += 1
                progress_bar.progress(completed / len(chunks))
        
        status_text.empty()
        progress_bar.empty()
        
        raw_df = pd.DataFrame(all_users)
        return _merge_duplicate_users(raw_df, pass_prefix=pass_prefix)

    except Exception as e:
        st.error(f"AI Extraction Error: {str(e)[:200]}")
        return None

def apply_ai_smart_context(df, command, api_key):
    """
    Applies natural language commands to the user dataframe using AI.
    Includes: 60s timeout, exponential-backoff retry, column allowlist guard.
    """
    api_key_str = str(api_key).strip()
    if api_key_str.startswith("AIzaSy"):
        client = OpenAI(
            api_key=api_key_str,
            base_url="https://generativelanguage.googleapis.com/v1beta/openai/"
        )
        model = "gemini-2.5-flash"
    else:
        client = OpenAI(api_key=api_key_str)
        model = "gpt-4o-mini"

    # Context size guard: truncate rows if JSON would be too large
    context_df = df
    context_json = context_df.to_json(orient='records')
    if len(context_json) > _MAX_AI_CONTEXT_KB * 1024:
        # Send only first 200 rows and warn
        context_df = df.head(200)
        context_json = context_df.to_json(orient='records')
        st.warning("⚠️ Dataset is large — AI command applied to first 200 rows only.")

    prompt = f"""
    You are a User Master Data Expert. The user wants to modify the following staff list.

    USER COMMAND: {command}

    CURRENT DATA (JSON):
    {context_json}

    INSTRUCTIONS:
    1. Understand the user's request (e.g., "Set role to X for all users in Y").
    2. Respond ONLY with a valid JSON list of objects representing the UPDATED rows.
    3. Each object MUST contain the '#' (serial number) to identify the row and ONLY the fields that changed.
    4. If the user wants to "Delete" or "Remove", return a field "::action": "delete" for those rows.

    FORMAT EXAMPLE:
    [{{"#": 1, "roles": "Admin|Doctor"}}, {{"#": 5, "isEnabled": "No"}}]
    """

    last_error = "Unknown error"
    for attempt in range(_AI_RETRY_ATTEMPTS):
        try:
            completion = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are a data patching engine. Return ONLY JSON."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"},
                timeout=60,
            )

            raw_res = completion.choices[0].message.content
            if not raw_res or not raw_res.strip():
                last_error = "AI returned an empty response."
                time.sleep(_AI_RETRY_BASE_WAIT * (2 ** attempt))
                continue

            res_data = json.loads(raw_res)

            # Normalise: AI may return {"updates": [...]} or bare list
            updates = res_data.get('updates', res_data) if isinstance(res_data, dict) else res_data
            if not isinstance(updates, list):
                last_error = "AI returned invalid format (expected a JSON list)."
                break

            # Apply updates with column allowlist guard
            new_df = df.copy()
            affected = 0
            blocked_fields = set()
            
            # Safety: Prevent mass-hallucination (If AI tries to change > 80% of rows)
            if len(updates) > len(df) * 0.8 and len(df) > 10:
                 return None, "⚠️ AI suggested a mass-change of >80% of your data. Operation blocked for safety."

            for update in updates:
                if '#' not in update:
                    continue
                idx = new_df[new_df['#'] == update['#']].index
                if idx.empty:
                    continue
                row_i = idx[0]
                for field, val in update.items():
                    if field == '#':
                        continue
                    if field not in _AI_ALLOWED_EDIT_COLS:
                        blocked_fields.add(field)   # log but don't apply
                        continue
                    if field in new_df.columns:
                        new_df.at[row_i, field] = val
                        affected += 1

            if blocked_fields:
                print(f"[AI Safety] Blocked hallucinated field(s): {blocked_fields}")

            summary = f"AI applied changes to {len(updates)} row(s) ({affected} cell update(s))."
            if blocked_fields:
                summary += f" ⚠️ Blocked invalid field(s): {', '.join(sorted(blocked_fields))}."
            return new_df, summary

        except json.JSONDecodeError as e:
            last_error = f"AI returned malformed JSON: {e}"
            print(f"[Smart Context attempt {attempt+1}] JSONDecodeError: {e}")
            break   # No point retrying a parse error

        except Exception as e:
            err_str = str(e)
            last_error = err_str[:200]
            print(f"[Smart Context attempt {attempt+1}] {type(e).__name__}: {err_str}")
            # Rate limit or transient: back off and retry
            if "429" in err_str or "timeout" in err_str.lower() or "connection" in err_str.lower():
                wait = _AI_RETRY_BASE_WAIT * (2 ** attempt)
                print(f"  Retrying in {wait}s...")
                time.sleep(wait)
                continue
            break  # Non-retryable error

    return None, f"AI command failed after {_AI_RETRY_ATTEMPTS} attempt(s): {last_error}"
def validate_master_data(df: pd.DataFrame):
    """
    Performance-optimized validation service for User Master Data.
    Uses vectorized pandas masks to handle large datasets efficiently.
    
    Returns: (errors_list, warnings_list)
    """
    errors = []
    warnings = []
    
    if df.empty:
        return errors, warnings

    # Regex patterns
    email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    mobile_regex = r'^\+?[0-9\s-]{8,20}$'

    # 1. Vectorized Mandatory Check (userName)
    # Identifies rows where userName is null, empty, or placeholder strings
    unames = df['userName'].astype(str).str.strip().str.lower()
    mask_missing_uname = (unames == '') | (unames == 'nan') | (unames == 'none') | (unames == '-')
    
    if mask_missing_uname.any():
        bad_indices = df[mask_missing_uname]['#'].tolist()
        for b_id in bad_indices:
            errors.append(f"Row {b_id}: Missing mandatory **userName**")

    # 2. Vectorized Email Check (Only for non-empty values)
    if 'email' in df.columns:
        emails = df['email'].astype(str).str.strip()
        # Create mask for rows that ARE NOT empty but DO NOT match regex
        has_email = (emails != '') & (emails != 'nan') & (emails != 'none') & (emails != '-')
        invalid_email = has_email & (~emails.str.match(email_regex, na=False))
        
        if invalid_email.any():
            bad_rows = df[invalid_email][['#', 'email']].values
            for b_id, b_val in bad_rows:
                warnings.append(f"Row {b_id}: Invalid **email** format ('{b_val}')")

    # 3. Vectorized Mobile Check (Only for non-empty values)
    if 'mobile' in df.columns:
        mobiles = df['mobile'].astype(str).str.strip()
        has_mobile = (mobiles != '') & (mobiles != 'nan') & (mobiles != 'none') & (mobiles != '-')
        invalid_mobile = has_mobile & (~mobiles.str.match(mobile_regex, na=False))
        
        if invalid_mobile.any():
            bad_rows = df[invalid_mobile][['#', 'mobile']].values
            for b_id, b_val in bad_rows:
                warnings.append(f"Row {b_id}: Invalid **mobile** format ('{b_val}')")

    return errors, warnings
